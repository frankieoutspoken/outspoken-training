import json
from typing import Any
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from claude_agent_sdk import tool
from src.config import OUTPUT_DIR


@tool(
    "generate_beauty_profile",
    "Generate a branded Midnight Cosmetics beauty profile document (.docx). Provide customer name, skin type, preferred look, concerns, and product recommendations as a JSON object.",
    {"profile_data": str},
)
async def generate_beauty_profile(args: dict[str, Any]) -> dict[str, Any]:
    try:
        data = json.loads(args.get("profile_data", "{}"))
    except json.JSONDecodeError:
        return {
            "content": [{"type": "text", "text": "Invalid profile data. Provide valid JSON."}],
            "is_error": True,
        }

    name = data.get("name", "Customer")
    skin_type = data.get("skin_type", "Not specified")
    preferred_look = data.get("preferred_look", "Not specified")
    concerns = data.get("concerns", [])
    recommendations = data.get("recommendations", [])
    morning_routine = data.get("morning_routine", "")
    evening_routine = data.get("evening_routine", "")
    tips = data.get("tips", "")

    doc = Document()

    # Title
    title = doc.add_heading("Midnight Cosmetics", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(75, 0, 130)

    subtitle = doc.add_heading("Your Personalized Beauty Profile", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Prepared for: {name}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Skin Profile
    doc.add_heading("Your Skin Profile", level=2)
    doc.add_paragraph(f"Skin Type: {skin_type}")
    doc.add_paragraph(f"Preferred Look: {preferred_look}")
    if concerns:
        doc.add_paragraph(f"Key Concerns: {', '.join(concerns) if isinstance(concerns, list) else concerns}")

    # Recommendations
    if recommendations:
        doc.add_heading("Recommended Products", level=2)
        for rec in recommendations:
            if isinstance(rec, dict):
                p = doc.add_paragraph()
                run = p.add_run(rec.get("name", "Product"))
                run.bold = True
                if rec.get("price"):
                    p.add_run(f" — ${rec['price']}")
                if rec.get("reason"):
                    doc.add_paragraph(rec["reason"], style="List Bullet")
            else:
                doc.add_paragraph(str(rec), style="List Bullet")

    # Routines
    if morning_routine:
        doc.add_heading("Morning Routine", level=2)
        doc.add_paragraph(morning_routine)

    if evening_routine:
        doc.add_heading("Evening Routine", level=2)
        doc.add_paragraph(evening_routine)

    # Tips
    if tips:
        doc.add_heading("Application Tips", level=2)
        doc.add_paragraph(tips)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("Midnight Cosmetics — Beauty, Personalized.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    filename = f"beauty_profile_{name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(filepath)

    return {"content": [{"type": "text", "text": f"Beauty profile generated: {filepath}"}]}
